# gerar_escalas.py
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from tkcalendar import Calendar
from docx import Document
from docx.shared import Pt
import os
import platform
from copy import deepcopy
from historico_manager import carregar_historico, salvar_historico

class GerarEscalas:
    def __init__(self, root, turmas, alunos_por_turma):
        self.root = root
        self.root.title("Gerar Escalas")
        self.root.geometry("900x700")

        self.turmas = turmas
        self.alunos_por_turma = alunos_por_turma
        self.dias_selecionados = []
        self.arquivos_gerados = []
        self.historico = carregar_historico()

        # Estilo personalizado
        self.style = ttk.Style()
        self.style.configure("Custom.TButton",
                             font=("Helvetica", 12, "bold"),
                             padding=10)
        self.style.configure("Custom.TLabel",
                             font=("Helvetica", 16),
                             foreground="#333333")
        self.style.configure("Custom.TCheckbutton",
                             font=("Helvetica", 12))

        # Interface gráfica
        self.create_widgets()

    def create_widgets(self):
        # Frame principal
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Título
        label_title = ttk.Label(main_frame, text="Gerar Escalas", style="Custom.TLabel")
        label_title.pack(pady=10)

        # Frame das turmas
        frame_turmas = ttk.LabelFrame(main_frame, text="Selecione as Turmas", padding="10")
        frame_turmas.pack(fill=tk.BOTH, expand=True, pady=10)

        # Canvas com scrollbar para as turmas
        canvas = tk.Canvas(frame_turmas)
        scrollbar = ttk.Scrollbar(frame_turmas, orient="vertical", command=canvas.yview)
        self.turmas_frame = ttk.Frame(canvas)

        self.turmas_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(
                scrollregion=canvas.bbox("all")
            )
        )

        canvas.create_window((0, 0), window=self.turmas_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Lista de turmas com caixas de seleção organizadas em colunas
        self.turma_vars = {}
        colunas = 4  # Número de colunas
        for index, turma in enumerate(self.turmas):
            var = tk.BooleanVar()
            chk = ttk.Checkbutton(self.turmas_frame, text=turma, variable=var, style="Custom.TCheckbutton")
            col = index % colunas
            row = index // colunas
            chk.grid(row=row, column=col, sticky='w', padx=10, pady=5)
            self.turma_vars[turma] = var

        # Ajuste das colunas no turmas_frame
        for col in range(colunas):
            self.turmas_frame.columnconfigure(col, weight=1)

        # Frame dos botões
        frame_botoes = ttk.Frame(main_frame)
        frame_botoes.pack(pady=20)

        # Botões
        btn_definir_dias = ttk.Button(frame_botoes, text="Definir Dias", command=self.definir_dias, style="Custom.TButton")
        btn_definir_dias.grid(row=0, column=0, padx=10, pady=10)

        btn_gerar_escala = ttk.Button(frame_botoes, text="Gerar Escala", command=self.gerar_escala, style="Custom.TButton")
        btn_gerar_escala.grid(row=0, column=1, padx=10, pady=10)

        btn_proxima_escala = ttk.Button(frame_botoes, text="Próxima Escala", command=self.proxima_escala, style="Custom.TButton")
        btn_proxima_escala.grid(row=0, column=2, padx=10, pady=10)

        btn_visualizar_escala = ttk.Button(frame_botoes, text="Visualizar Escala", command=self.visualizar_escala, style="Custom.TButton")
        btn_visualizar_escala.grid(row=0, column=3, padx=10, pady=10)

        btn_limpar_dados = ttk.Button(frame_botoes, text="Limpar Dados", command=self.limpar_dados, style="Custom.TButton")
        btn_limpar_dados.grid(row=0, column=4, padx=10, pady=10)

    def definir_dias(self):
        self.calendar_window = tk.Toplevel(self.root)
        self.calendar_window.title("Selecionar Dias")
        self.calendar_window.geometry("400x400")
        self.calendar = Calendar(self.calendar_window, selectmode='day', date_pattern='yyyy-mm-dd')
        self.calendar.pack(padx=10, pady=10)

        label_instrucao = ttk.Label(self.calendar_window, text="Selecione os dias clicando nas datas desejadas.")
        label_instrucao.pack(pady=5)

        btn_confirmar_dias = ttk.Button(self.calendar_window, text="Confirmar Dias", command=self.obter_dias_selecionados, style="Custom.TButton")
        btn_confirmar_dias.pack(pady=10)

        self.calendar.bind("<<CalendarSelected>>", self.on_date_selected)

    def on_date_selected(self, event):
        date = self.calendar.get_date()
        if date in self.dias_selecionados:
            self.dias_selecionados.remove(date)
            self.calendar.calevent_remove('highlight', date)
        else:
            self.dias_selecionados.append(date)
            self.calendar.calevent_create(date, 'Selecionado', 'highlight')
            self.calendar.tag_config('highlight', background='lightblue', foreground='black')

        self.calendar.selection_clear()

    def obter_dias_selecionados(self):
        if not self.dias_selecionados:
            messagebox.showwarning("Aviso", "Nenhum dia selecionado.")
        else:
            messagebox.showinfo("Dias Selecionados", f"Dias selecionados: {', '.join(self.dias_selecionados)}")
            self.calendar_window.destroy()

    def gerar_escala(self):
        turmas_selecionadas = [turma for turma, var in self.turma_vars.items() if var.get()]
        if not turmas_selecionadas:
            messagebox.showwarning("Aviso", "Nenhuma turma selecionada.")
            return
        if not self.dias_selecionados:
            messagebox.showwarning("Aviso", "Nenhum dia selecionado.")
            return

        modelo_path = filedialog.askopenfilename(title="Selecione o Modelo de Documento",
                                                 filetypes=[("Documentos Word", "*.docx")])
        if not modelo_path:
            messagebox.showwarning("Aviso", "Nenhum modelo selecionado.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Documentos Word", "*.docx")])
        if file_path:
            self.criar_documento_unico(modelo_path, file_path, turmas_selecionadas)
            messagebox.showinfo("Sucesso", f"Escala gerada com sucesso em {file_path}.")
            self.arquivos_gerados.append(file_path)

    def criar_documento_unico(self, modelo_path, file_path, turmas_selecionadas):
        documento = Document(modelo_path)

        turma_counter = 1  # Contador para diferenciar os placeholders de cada turma
        substitutions = {}

        for turma in turmas_selecionadas:
            alunos = self.alunos_por_turma.get(turma, [])
            if alunos:
                alunos_ordenados = sorted(alunos)
                chefe = alunos_ordenados[0]
                subchefe = alunos_ordenados[1] if len(alunos_ordenados) > 1 else ''
                limpeza_alunos = sorted(alunos, reverse=True)[:3]
            else:
                chefe = ''
                subchefe = ''
                limpeza_alunos = [''] * 3

            # Criar os placeholders dinamicamente
            substitutions.update({
                f'{{{{turma{turma_counter}}}}}': turma,
                f'{{{{chefe{turma_counter}}}}}': chefe,
                f'{{{{subchefe{turma_counter}}}}}': subchefe,
                f'{{{{aluno{turma_counter}_1}}}}': limpeza_alunos[0] if len(limpeza_alunos) > 0 else '',
                f'{{{{aluno{turma_counter}_2}}}}': limpeza_alunos[1] if len(limpeza_alunos) > 1 else '',
                f'{{{{aluno{turma_counter}_3}}}}': limpeza_alunos[2] if len(limpeza_alunos) > 2 else '',
            })

            turma_counter += 1
        # Após criar as substitutions
        print("Substitutions:", substitutions)
    

        # Substituir os placeholders no documento
        self.substituir_placeholders(documento, substitutions)
        # Substituir o placeholder de data, se existir
        data_substitutions = {'{{data}}': ', '.join(self.dias_selecionados)}
        self.substituir_placeholders(documento, data_substitutions)

        # Salvar o documento final
        documento.save(file_path)
        self.arquivos_gerados.append(file_path)

    def proxima_escala(self):
        turmas_selecionadas = [turma for turma, var in self.turma_vars.items() if var.get()]
        if not turmas_selecionadas:
            messagebox.showwarning("Aviso", "Nenhuma turma selecionada.")
            return
        if not self.dias_selecionados:
            messagebox.showwarning("Aviso", "Nenhum dia selecionado.")
            return

        modelo_path = filedialog.askopenfilename(
            title="Selecione o Modelo de Documento", filetypes=[("Documentos Word", "*.docx")]
        )
        if not modelo_path:
            messagebox.showwarning("Aviso", "Nenhum modelo selecionado.")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx", filetypes=[("Documentos Word", "*.docx")]
        )
        if file_path:
            self.criar_proxima_escala_documento_unico(modelo_path, file_path, turmas_selecionadas)
            messagebox.showinfo("Sucesso", f"Próxima escala gerada com sucesso em {file_path}.")
            self.arquivos_gerados.append(file_path)

    def criar_proxima_escala_documento_unico(self, modelo_path, file_path, turmas_selecionadas):
        documento = Document(modelo_path)
        substitutions = {}
        turma_counter = 1

        for turma in turmas_selecionadas:
            alunos = self.alunos_por_turma.get(turma, [])
            if not alunos:
                messagebox.showwarning("Aviso", f"A turma {turma} não possui alunos cadastrados.")
                continue

            historico_turma = self.historico.get(turma, {
                "indices_utilizados": [],
                "proximo_indice": 0
            })

            alunos_ordenados = sorted(alunos)
            total_alunos = len(alunos_ordenados)

            if len(historico_turma["indices_utilizados"]) >= total_alunos:
                historico_turma["indices_utilizados"] = []
                historico_turma["proximo_indice"] = 0

            indices_selecionados = []
            chefe_indice = self.obter_proximo_indice(historico_turma, total_alunos)
            indices_selecionados.append(chefe_indice)
            chefe = alunos_ordenados[chefe_indice]

            subchefe_indice = self.obter_proximo_indice(historico_turma, total_alunos)
            indices_selecionados.append(subchefe_indice)
            subchefe = alunos_ordenados[subchefe_indice]

            limpeza_alunos = []
            for _ in range(3):
                indice_aluno = self.obter_proximo_indice(historico_turma, total_alunos)
                indices_selecionados.append(indice_aluno)
                limpeza_alunos.append(alunos_ordenados[indice_aluno])

            self.historico[turma] = historico_turma
            salvar_historico(self.historico)

            # Criar os placeholders dinamicamente
            substitutions.update({
                f'{{{{turma{turma_counter}}}}}': turma,
                f'{{{{chefe{turma_counter}}}}}': chefe,
                f'{{{{subchefe{turma_counter}}}}}': subchefe,
                f'{{{{aluno{turma_counter}_1}}}}': limpeza_alunos[0],
                f'{{{{aluno{turma_counter}_2}}}}': limpeza_alunos[1],
                f'{{{{aluno{turma_counter}_3}}}}': limpeza_alunos[2],
            })

            turma_counter += 1

        # Substituir os placeholders no documento
        self.substituir_placeholders(documento, substitutions)
        # Substituir o placeholder de data, se existir
        data_substitutions = {'{{data}}': ', '.join(self.dias_selecionados)}
        self.substituir_placeholders(documento, data_substitutions)

        # Salvar o documento final
        documento.save(file_path)
        self.arquivos_gerados.append(file_path)

    def obter_proximo_indice(self, historico_turma, total_alunos):
        indice = historico_turma["proximo_indice"] % total_alunos
        while indice in historico_turma["indices_utilizados"]:
            historico_turma["proximo_indice"] += 1
            indice = historico_turma["proximo_indice"] % total_alunos
        historico_turma["indices_utilizados"].append(indice)
        historico_turma["proximo_indice"] += 1
        return indice

    def substituir_placeholders(self, elemento, substitutions):
        if hasattr(elemento, 'paragraphs'):
            for paragraph in elemento.paragraphs:
                self.substituir_placeholders_no_paragrafo(paragraph, substitutions)

        if hasattr(elemento, 'tables'):
            for table in elemento.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self.substituir_placeholders(cell, substitutions)

    def substituir_placeholders_no_paragrafo(self, paragraph, substitutions):
        for key, value in substitutions.items():
            if key in paragraph.text:
                inline = paragraph.runs
                for run in inline:
                    if key in run.text:
                        run.text = run.text.replace(key, value)

    def visualizar_escala(self):
        if not self.arquivos_gerados:
            messagebox.showwarning("Aviso", "Nenhuma escala gerada para visualizar.")
            return

        if len(self.arquivos_gerados) > 1:
            arquivo_selecionado = filedialog.askopenfilename(title="Selecione a Escala para Visualizar",
                                                             initialdir=os.path.dirname(self.arquivos_gerados[-1]),
                                                             filetypes=[("Documentos Word", "*.docx")])
            if not arquivo_selecionado:
                return
        else:
            arquivo_selecionado = self.arquivos_gerados[-1]

        try:
            if platform.system() == 'Windows':
                os.startfile(arquivo_selecionado)
            elif platform.system() == 'Darwin':
                os.system(f'open "{arquivo_selecionado}"')
            else:
                os.system(f'xdg-open "{arquivo_selecionado}"')
        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível abrir o arquivo.\nDetalhes: {e}")

    def limpar_dados(self):
        resposta = messagebox.askyesno("Confirmação", "Tem certeza de que deseja limpar todos os dados das escalas?")
        if resposta:
            try:
                # Remove o arquivo de histórico
                if os.path.exists('historico_escalas.json'):
                    os.remove('historico_escalas.json')
                # Limpa o histórico carregado
                self.historico = {}
                messagebox.showinfo("Sucesso", "Todos os dados das escalas foram limpos com sucesso.")
            except Exception as e:
                messagebox.showerror("Erro", f"Não foi possível limpar os dados.\nDetalhes: {e}")
